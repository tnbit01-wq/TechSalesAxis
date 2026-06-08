"""
JD Parser Service - Extracts and parses Job Description documents
Supports PDF, DOCX, DOC, and TXT files
Uses AI to extract structured job fields from raw text
"""

import io
import re
import logging
from typing import Dict, Any, Optional

logger = logging.getLogger(__name__)

# Maximum file size: 10MB
MAX_FILE_SIZE = 10 * 1024 * 1024
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt"}


class JDParserService:
    """
    Parses Job Description files into structured job form data.
    Extracts text from PDF/DOCX/TXT, then uses AI to structure it.
    """

    @staticmethod
    def extract_text_from_pdf(file_bytes: bytes) -> str:
        """Extract text from a PDF file."""
        try:
            try:
                from pypdf import PdfReader
            except ImportError:
                from PyPDF2 import PdfReader
            
            reader = PdfReader(io.BytesIO(file_bytes))
            text_parts = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            return "\n".join(text_parts)
        except ImportError:
            logger.error("Neither pypdf nor PyPDF2 is installed. Install with: pip install pypdf")
            raise ValueError("PDF parsing is not available. Server missing pypdf dependency.")
        except Exception as e:
            logger.error(f"PDF extraction failed: {str(e)}")
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")

    @staticmethod
    def extract_text_from_docx(file_bytes: bytes) -> str:
        """Extract text from a DOCX file."""
        try:
            from docx import Document
            doc = Document(io.BytesIO(file_bytes))
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text.strip())
            return "\n".join(text_parts)
        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            raise ValueError("DOCX parsing is not available. Server missing python-docx dependency.")
        except Exception as e:
            logger.error(f"DOCX extraction failed: {str(e)}")
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")

    @staticmethod
    def extract_text_from_txt(file_bytes: bytes) -> str:
        """Extract text from a plain text file."""
        try:
            # Try UTF-8 first, then fall back to latin-1
            try:
                return file_bytes.decode("utf-8")
            except UnicodeDecodeError:
                return file_bytes.decode("latin-1")
        except Exception as e:
            logger.error(f"TXT extraction failed: {str(e)}")
            raise ValueError(f"Failed to read text file: {str(e)}")

    @classmethod
    def extract_text(cls, file_bytes: bytes, filename: str) -> str:
        """Extract text from a file based on its extension."""
        ext = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        
        if ext not in SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type: {ext}. "
                f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
            )

        if len(file_bytes) > MAX_FILE_SIZE:
            raise ValueError(f"File too large. Maximum size is {MAX_FILE_SIZE // (1024*1024)}MB.")

        if len(file_bytes) == 0:
            raise ValueError("File is empty.")

        if ext == ".pdf":
            return cls.extract_text_from_pdf(file_bytes)
        elif ext in (".docx", ".doc"):
            return cls.extract_text_from_docx(file_bytes)
        elif ext == ".txt":
            return cls.extract_text_from_txt(file_bytes)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

    @staticmethod
    def build_parse_prompt(raw_text: str) -> str:
        """Build the AI prompt to parse JD text into structured fields."""
        # Truncate very long texts to stay within token limits
        truncated = raw_text[:8000] if len(raw_text) > 8000 else raw_text

        return f"""
You are an expert Job Description parser for IT Tech Sales roles.

Analyze the following Job Description text and extract structured information.

JD TEXT:
\"\"\"
{truncated}
\"\"\"

CRITICAL INSTRUCTIONS:
1. Extract the EXACT job title from the JD. If not found, infer from context.
2. Extract or compose a professional role description (2-3 paragraphs).
3. Extract key requirements as bullet points (5-8 items).
4. Extract technical and soft skills mentioned (top 5-8 skills).
5. Determine the experience band from any years-of-experience mentioned:
   - 0-1 years → "fresher"
   - 1-5 years → "mid"  
   - 5-10 years → "senior"
   - 10+ years → "leadership"
6. Determine job_type from the JD: "onsite", "remote", or "hybrid"
7. Extract location if mentioned.
8. Extract or infer salary range if mentioned. If not mentioned, put "Market Competitive".
9. Extract number_of_positions if mentioned, default to 1.

Return ONLY valid JSON with these exact keys:
{{
    "title": "string",
    "description": "string (2-3 paragraphs)",
    "requirements": ["requirement 1", "requirement 2", ...],
    "skills_required": ["skill 1", "skill 2", ...],
    "experience_band": "fresher|mid|senior|leadership",
    "job_type": "onsite|remote|hybrid",
    "location": "string or empty",
    "salary_range": "string",
    "number_of_positions": 1
}}
"""

    @classmethod
    async def parse_jd_file(cls, file_bytes: bytes, filename: str) -> Dict[str, Any]:
        """
        Full pipeline: extract text from file → AI parse → return structured data.
        Uses the recruiter_service AI helper for consistency.
        """
        from src.services.recruiter_service import recruiter_service

        # Step 1: Extract raw text
        raw_text = cls.extract_text(file_bytes, filename)

        if not raw_text or len(raw_text.strip()) < 50:
            raise ValueError(
                "Could not extract enough text from the file. "
                "Please ensure the document contains readable text."
            )

        logger.info(f"[JD PARSER] Extracted {len(raw_text)} chars from {filename}")

        # Step 2: Use AI to parse into structured fields
        prompt = cls.build_parse_prompt(raw_text)
        system_message = (
            "You are an elite Job Description parser. "
            "Extract structured hiring data from JD documents with high accuracy. "
            "Always return valid JSON."
        )

        ai_result = await recruiter_service._call_ai_json(prompt, system_message)

        if not ai_result or not ai_result.get("title"):
            # Fallback: try basic regex extraction
            logger.warning("[JD PARSER] AI parsing failed, using fallback extraction")
            return cls._fallback_extract(raw_text)

        # Normalize the result
        return {
            "title": ai_result.get("title", "Untitled Role"),
            "description": ai_result.get("description", raw_text[:500]),
            "requirements": ai_result.get("requirements") or [],
            "skills_required": ai_result.get("skills_required") or [],
            "experience_band": ai_result.get("experience_band", "mid"),
            "job_type": ai_result.get("job_type", "onsite"),
            "location": ai_result.get("location", ""),
            "salary_range": ai_result.get("salary_range", "Market Competitive"),
            "number_of_positions": ai_result.get("number_of_positions", 1),
            "is_ai_generated": True,
            "source": "jd_upload",
        }

    @staticmethod
    def _fallback_extract(text: str) -> Dict[str, Any]:
        """Basic regex-based extraction when AI fails."""
        # Try to find title (usually first meaningful line)
        lines = [l.strip() for l in text.split("\n") if l.strip()]
        title = lines[0] if lines else "Untitled Role"
        
        # Truncate title if too long
        if len(title) > 100:
            title = title[:97] + "..."

        # Experience detection
        exp_match = re.search(r'(\d+)\+?\s*(?:years?|yrs?)', text, re.IGNORECASE)
        years = int(exp_match.group(1)) if exp_match else 3
        if years < 2:
            band = "fresher"
        elif years < 5:
            band = "mid"
        elif years < 10:
            band = "senior"
        else:
            band = "leadership"

        # Location detection
        location = ""
        loc_match = re.search(
            r'(?:location|based\s+in|city)\s*[:\-]\s*([A-Za-z\s,]+?)(?:\n|$)',
            text, re.IGNORECASE
        )
        if loc_match:
            location = loc_match.group(1).strip()

        # Job type detection
        text_lower = text.lower()
        if "remote" in text_lower:
            job_type = "remote"
        elif "hybrid" in text_lower:
            job_type = "hybrid"
        else:
            job_type = "onsite"

        return {
            "title": title,
            "description": "\n".join(lines[:10]) if len(lines) > 1 else text[:500],
            "requirements": [],
            "skills_required": [],
            "experience_band": band,
            "job_type": job_type,
            "location": location,
            "salary_range": "Market Competitive",
            "number_of_positions": 1,
            "is_ai_generated": False,
            "source": "jd_upload_fallback",
        }
